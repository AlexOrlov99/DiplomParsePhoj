from datetime import datetime
import glob, re
from docx import Document
from docx.shared import Pt
from docx.opc.constants import RELATIONSHIP_TYPE as RT


class DocRezume():

    MIN_PARAGRAPHS_COUNT = 1
    NAME = 0
    SURNAME = 1
    PATRONYMIC = 2
    MIN_ITEMS_LEN = 2
    
    skills_pattern: list = [
        re.compile(r'[А-Яа-я]+\s?[Ии]?\s?Н?н?авыки[-:]?'),
        re.compile(r'квалификация[-:]?'),
        ]
    education_pattern = re.compile(r'образование[-:]?')
    experience_pattern = re.compile(r'(профессиональный)?'
                                    r'\s*о?(пыт)?\s*(работы)?[-:]?')

    def __init__(self):
        
        self.full_name:str = None
        self.email:str = None
        self.phone:str = None
        self.education = None
        self.experience = None
        self.skills = None
        self.file = None

    def __str__(self) -> str:
        
        str_value:str = '0:^10'.format(self.full_name)
        return str_value
   
    def find_full_name(self, items):
        full_name: str =''
        name: str = ''.join(items[self.NAME])
        surname: str = ''.join(items[self.SURNAME])
        if len(items) == self.PATRONYMIC + 1:
            full_name = name + ' ' + surname + \
                        ' '  + ''.join(items[self.PATRONYMIC])
        else:
            full_name = name + ' ' + surname
        return full_name

    def find_email(self, items):
        email_pattern = re.compile(r'\b[a-zA-Z0-9][a-zA-Z0-9._-]*'
                                    r'[^._-]@[a-z]{,6}\.[a-z]{2,3}')
        for item in items:
            checked_email: re.match = re.fullmatch(email_pattern, item)
            if checked_email:
                return item

    def find_phone(self, items):
        number_pattern = re.compile(r'\+\d{1,3}|\(\d{3}\)|\(\d\)|\d{2,3}')
        phone_pattern = re.compile(r'^((8|\+\d+)[\- ]?)?'
                                    r'(\(?\d{3}\)?[\- ]?)?[\d\- ]{7,10}$')
        number: str = ''
        for item in items:
            checked_phone: re.match = re.fullmatch(phone_pattern, item)
            if checked_phone:
                return item
            else:           
                checked_number: re.match = re.fullmatch(number_pattern, item)
                if checked_number and self.phone == None:
                    number = number + item
                checked_phone: re.match = re.fullmatch(phone_pattern, number)            
                if checked_phone:
                    return number

    def find_item_in_paragraphs(self, doc):

        search_word_skills = None
        search_word_education = None
        search_word_experience = None

        if len(doc.paragraphs) > self.MIN_PARAGRAPHS_COUNT:
            for paragraph in doc.paragraphs:
                items = paragraph.text.split()
                
                if len(items) <=0 or items[0] == 'Резюме':
                    continue
                
                if self.full_name == None and len(items)>= self.MIN_ITEMS_LEN:
                    self.full_name = self.find_full_name(items)
            
                if self.email == None:
                    self.email = self.find_email(items)

                if self.phone == None and len(items) >= self.MIN_ITEMS_LEN:
                    self.phone = self.find_phone(items)
                        
                if search_word_education == None:
                    search_word: str = ''.join(items[0])
                    checked_education: re.match = re.fullmatch(
                        self.education_pattern,search_word.lower())
                    search_word_education = search_word if checked_education \
                        else None

                if search_word_experience == None and len(items) > 1:
                    search_word = ''.join(items[0]) + ' ' + ''.join(items[-1:])
                    checked_experience: re.match  = re.fullmatch(
                        self.experience_pattern, 
                        search_word.lower())
                    search_word_experience = search_word \
                        if checked_experience else None

                if search_word_skills == None:
                    search_word = ''.join(items[0])
                    if len(items) > 1:
                        search_word = ''.join(items[0]) + ' ' + ''.join(items[-1:])
                    for i in range(len(self.skills_pattern)):
                        checked_skills: re.match = re.fullmatch(
                            self.skills_pattern[i],
                            search_word.lower()
                            )
                        search_word_skills = search_word if checked_skills \
                            else None
 
    def find_item_in_tables(self, doc):

        for table in doc.tables:
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    items = cell.text.split()

                    if self.full_name == None and \
                        len(items)>= self.MIN_ITEMS_LEN:
                        self.full_name = self.find_full_name(items)

                    if self.email == None:
                        self.email = self.find_email(items)
                            
                    if self.phone == None and \
                        len(items) >= self.MIN_ITEMS_LEN:
                        self.phone = self.find_phone(items)

                    if len(items) <=0 :
                        continue

                    if self.skills == None:
                        search_word = ''
                        if len(items) == 1:
                            search_word = ''.join(items[0])
                        elif len(items) > 1:
                            for item in items:
                                search_word += ''.join(item) + ' '
                            search_word = search_word[:len(search_word)-1]
                        for pattern in self.skills_pattern:
                            checked_skills = re.fullmatch(pattern, 
                                                        search_word.lower())
                            if checked_skills:
                                self.skills = table.cell(i, j + 1).text
                    
                    if self.education == None:
                        search_word: str  = ''.join(items[0])
                        checked_education = re.fullmatch(
                            self.education_pattern, search_word.lower())
                        if checked_education:
                            self.education = table.cell(i, j + 1).text

                    if self.experience == None and len(items) > 1:
                            search_word: str = ''.join(items[0]) + \
                                ' ' + ''.join(items[-1:])
                            checked_experience: re.match = re.fullmatch(
                                self.experience_pattern, search_word.lower())
                            self.experience = table.cell(i, j + 1).text \
                                if checked_experience else None
                                    
    def find_email_in_hyperlinks(self, doc):
        rels = doc.part.rels
        for rel in rels:
            if rels[rel].reltype == RT.HYPERLINK:
                link = rels[rel]._target.split(':')
                self.email = ''.join(link[-1:])
    
    def handle(self) -> None:
        """Handles data filling."""

        start: datetime = datetime.now()
        doc = Document(self.file)
        self.find_item_in_paragraphs(doc)
        self.find_item_in_tables(doc)
        self.find_email_in_hyperlinks(doc)
        print(
            "Parsing data: {} seconds".format(
                (datetime.now()-start).total_seconds()
            )
        )
    
# doc1 = 'Docs/Zadirov.docx'
# doc2 = 'Docs/Scoric.docx'
# doc3 = 'Docs/Zablodsky.docx'
# doc4 = 'Docs/Shihov.docx'
# doc5 = 'Docs/Ivanova.docx'
# doc6 = 'Docs/Moloda.docx'
# file_address = doc6

# rezume = DocRezume()
# rezume.handle(file_address)
# document = Document()
# style = document.styles['Normal']
# style.font.name = 'Arial'
# style.font.size = Pt(14)
# document.add_paragraph(rezume.full_name)
# document.add_paragraph(rezume.email)
# document.add_paragraph(rezume.phone)
# document.add_paragraph(rezume.education)
# document.add_paragraph(rezume.experience)
# document.add_paragraph(rezume.skills)
# document.save('test.docx')